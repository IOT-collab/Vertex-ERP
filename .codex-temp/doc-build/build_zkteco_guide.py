from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"D:\sanjeev\Web Forms\Vertex\new1\Vertex-ERP-Krishna_23.7\Vertex-ERP-Krishna_23.7")
OUT = ROOT / "Vertex ERP" / "Documentation" / "ZKTeco_K40_Pro_Integration_Guide.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK = "17365D"
LIGHT = "E8EEF5"
PALE = "F3F7FB"
GREEN = "2E7D32"
AMBER = "9C6500"
RED = "B42318"
GRAY = "5B6573"
WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
]:
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

for name, size, color, bold in [
    ("Guide Title", 27, DARK, True),
    ("Guide Subtitle", 13, GRAY, False),
    ("Step Title", 11, DARK, True),
    ("Small Label", 9, GRAY, True),
]:
    if name not in styles:
        s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    else:
        s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_after = Pt(5)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), "9360"); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[i])); tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[i] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)

def add_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, value in enumerate(headers):
        shade(hdr.cells[i], LIGHT)
        p = hdr.cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value); r.bold = True; r.font.color.rgb = RGBColor.from_string(DARK)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2 == 1: shade(cells[i], "F8FAFC")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value)); r.font.size = Pt(9.3)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

def add_callout(label, text, fill=PALE, color=DARK):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.cell(0,0); shade(c, fill); cell_margins(c, 140, 180, 140, 180)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  "); r.bold = True; r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)
    set_table_geometry(t, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p

def add_number(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p

def add_step(title, body):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title); r.bold = True; r.font.color.rgb = RGBColor.from_string(DARK)
    p2 = doc.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(8)

# Header / footer
header = sec.header.paragraphs[0]
header.text = "VERTEX ERP  |  BIOMETRIC INTEGRATION GUIDE"
header.style = styles["Small Label"]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("ZKTeco K40 Pro / ADMS Push  |  Internal Implementation Guide  |  ")
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
footer._p.append(fld)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(75)
p.paragraph_format.space_after = Pt(12)
r = p.add_run("ZKTeco K40 Pro")
r.font.name = "Calibri"; r.font.size = Pt(16); r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph(style="Guide Title")
p.add_run("Biometric Attendance\nIntegration Guide")
p = doc.add_paragraph(style="Guide Subtitle")
p.add_run("ASP.NET Core MVC (.NET 8) + PostgreSQL + ZKTeco ADMS Push")
doc.add_paragraph()
add_callout("Purpose", "Configure the K40 Pro, receive every punch in Vertex ERP, map device users to employees, and display daily IN/OUT attendance on the existing Attendance UI.")
add_table(["Document", "Value"], [
    ("Audience", "ERP administrator, implementation engineer, IT support"),
    ("Device", "ZKTeco K40 Pro, firmware 8.0.4.3, ADMS/Push mode"),
    ("Application", "Vertex ERP - ASP.NET Core MVC (.NET 8)"),
    ("Database", "PostgreSQL with Entity Framework Core"),
    ("Test topology", "Direct Ethernet: laptop 192.168.10.1, device 192.168.10.2"),
], [2400, 6960])
doc.add_paragraph("Version 1.0 | Prepared for implementation and support", style="Small Label")
doc.add_page_break()

doc.add_heading("1. Integration at a glance", level=1)
doc.add_paragraph("The K40 Pro uses ADMS Push. The device initiates HTTP requests to the ERP server. The ERP does not continuously poll the biometric machine.")
add_callout("Data flow", "Fingerprint punch -> K40 Pro -> HTTP port 8082 -> /iclock/cdata -> raw AttendanceLogs -> employee mapping -> daily attendance -> existing Attendance Management UI.", fill="EAF4FF")
add_table(["Layer", "Responsibility"], [
    ("K40 Pro", "Captures fingerprint and pushes raw punch records."),
    ("Network", "Provides a reachable server IP and allows inbound TCP 8082."),
    ("ADMS API", "Accepts device heartbeat, options, user data and ATTLOG payloads."),
    ("Raw storage", "Stores every received punch without discarding the original payload."),
    ("Mapping", "Connects the device user number to the existing Employee record."),
    ("Processing", "First punch = IN; last punch = OUT; difference = working hours."),
    ("Attendance UI", "Shows mapped daily attendance in the current Vertex ERP page."),
], [2200, 7160])

doc.add_heading("2. Prerequisites", level=1)
add_bullet("K40 Pro is powered on and the employee fingerprint is enrolled.")
add_bullet("Laptop/server and device have working Ethernet or Wi-Fi connectivity.")
add_bullet("Vertex ERP is configured to listen on HTTP port 8082 for the device.")
add_bullet("PostgreSQL is reachable and biometric migrations are applied.")
add_bullet("Windows Firewall allows inbound TCP 8082 on the active network profile.")
add_bullet("The device serial number is registered in Vertex ERP exactly as reported by the machine.")
add_callout("Important", "The K40 Pro screen may also show TCP COMM Port 4370. That port belongs to the older SDK/pull protocol. ADMS Push in this integration uses the Cloud Server port 8082.", fill="FFF4E5", color=AMBER)

doc.add_heading("3. Network setup", level=1)
doc.add_heading("3.1 Direct Ethernet setup (current test setup)", level=2)
add_table(["Setting", "Laptop Ethernet", "K40 Pro Ethernet"], [
    ("IP address", "192.168.10.1", "192.168.10.2"),
    ("Subnet mask", "255.255.255.0", "255.255.255.0"),
    ("Gateway", "Leave blank or 192.168.10.1", "192.168.10.1"),
    ("DNS", "Leave blank", "192.168.10.1"),
    ("DHCP", "Not applicable", "OFF"),
], [2400, 3480, 3480])
add_step("Step 1 - Configure the laptop", "Open Ethernet adapter IPv4 properties and set the static address 192.168.10.1 with subnet mask 255.255.255.0.")
add_step("Step 2 - Configure the machine", "On the K40 Pro, open Communication > Ethernet. Set IP address 192.168.10.2, subnet mask 255.255.255.0, gateway 192.168.10.1, DNS 192.168.10.1 and DHCP OFF.")
add_step("Step 3 - Verify the cable link", "Confirm Ethernet link lights are active, then ping 192.168.10.2 from the laptop. A reply confirms basic connectivity.")

doc.add_heading("3.2 Router/LAN setup (production option)", level=2)
doc.add_paragraph("Connect both systems to the same router. Prefer DHCP reservation or static addresses outside the router's automatic allocation range. The server address configured in ADMS must be the laptop/server LAN IP, not localhost and not the device IP.")
add_callout("Do not use", "127.0.0.1, localhost, 0.0.0.0, or a public IP that does not route back to the ERP server from the machine.", fill="FDECEC", color=RED)

doc.add_heading("4. Configure ADMS on the K40 Pro", level=1)
doc.add_paragraph("Open Menu > Comm. > Cloud Server Setting and enter the following values.")
add_table(["Machine field", "Required value", "Meaning"], [
    ("Server Mode", "ADMS", "Enables the device push protocol."),
    ("Enable Domain Name", "OFF", "Use a direct IP during local testing."),
    ("Server Address", "192.168.10.1", "IP of the laptop/server running Vertex ERP."),
    ("Server Port", "8082", "HTTP listener used by the ERP ADMS endpoint."),
    ("Enable Proxy Server", "OFF", "No proxy is required on a local network."),
    ("HTTPS", "OFF", "Current local listener is plain HTTP."),
], [2450, 2400, 4510])
add_number("Save the settings.")
add_number("Restart the K40 Pro.")
add_number("Start Vertex ERP before testing a punch.")
add_number("Wait up to one minute for the first heartbeat, then check Device Status and Last Sync Time.")

doc.add_heading("5. Configure Vertex ERP", level=1)
doc.add_heading("5.1 HTTP listener and endpoint", level=2)
doc.add_paragraph("The application must listen on all local interfaces so that the physical device can reach it. The configured URL is:")
add_callout("Listener", "http://0.0.0.0:8082", fill="EAF4FF")
add_table(["Method", "Endpoint", "Use"], [
    ("GET/POST", "/iclock/cdata", "Device registration, heartbeat and raw attendance uploads."),
    ("GET", "/iclock/getrequest", "Device command polling."),
    ("POST", "/iclock/devicecmd", "Device command result acknowledgement."),
    ("GET", "/BiometricDevices", "Device list, status, test connection and mappings."),
    ("GET", "/Main/Attendence", "Existing attendance page populated from biometric data."),
], [1300, 3000, 5060])
add_callout("HTTP redirect rule", "Exclude /iclock requests from automatic HTTPS redirection while the machine's HTTPS option is OFF. Otherwise the machine may receive a redirect it cannot complete.", fill="FFF4E5", color=AMBER)

doc.add_heading("5.2 Firewall", level=2)
doc.add_paragraph("Create an inbound Windows Firewall rule allowing TCP port 8082 for the active network profile. Limit the scope to the local subnet or device address in production.")
add_callout("Quick check", "From another computer on the same network, open http://SERVER-IP:8082/iclock/cdata?SN=TEST. Any HTTP response proves the port is reachable; a timeout means the listener or firewall still blocks it.")

doc.add_heading("5.3 Database objects", level=2)
add_table(["Table", "Purpose"], [
    ("BiometricDevices", "Device identity, serial number, model, network settings, active status and last sync time."),
    ("AttendanceLogs", "Immutable raw punch records, receive time, device user ID, punch time, verification type and raw payload."),
    ("EmployeeDeviceMapping", "Maps one device user ID to an existing Employee and supports multiple devices."),
], [2750, 6610])
doc.add_paragraph("Apply the EF Core biometric migration once per database. Do not delete raw attendance logs when reprocessing daily attendance.")

doc.add_heading("6. Register the biometric device", level=1)
add_number("Sign in to Vertex ERP as an administrator.")
add_number("Open Biometric Devices.")
add_number("Select Add Device.")
add_number("Enter a friendly name, model K40 Pro, the exact device serial number, branch (if applicable), server/network values, and mark the device Active.")
add_number("Save, then restart the K40 Pro or wait for its next heartbeat.")
add_number("Confirm Device Status becomes Online and Last Sync Time updates.")
add_callout("Serial matching", "ADMS identifies the device with the SN query value. The serial in BiometricDevices must match exactly; spaces and typing mistakes will prevent association.", fill="FFF4E5", color=AMBER)

doc.add_heading("7. Map device users to employees", level=1)
doc.add_paragraph("A punch can be stored successfully but still not appear under an employee until its Device User ID is mapped.")
add_number("Make one test punch on the K40 Pro.")
add_number("Open Biometric Devices > Employee Mappings.")
add_number("Locate the unmapped Device User ID received from the machine.")
add_number("Select the corresponding employee from the existing Employee Management module.")
add_number("Save the mapping and run attendance processing/backfill for the required date range.")
add_callout("Example", "If the machine sends user ID 121 and that fingerprint belongs to Sanjeev, map Device User ID 121 to Sanjeev's Employee record. The fingerprint template itself is not copied into AttendanceLogs.", fill="EAF4FF")

doc.add_heading("8. Attendance processing rules", level=1)
add_table(["Rule", "Result"], [
    ("One punch in a day", "IN time is set; OUT time remains blank until another punch arrives."),
    ("Two or more punches", "Earliest valid punch = IN; latest valid punch = OUT."),
    ("Working hours", "OUT time minus IN time."),
    ("Duplicate delivery", "Ignore an already stored device/event identity while preserving the original raw record."),
    ("Unmapped user", "Keep the raw log and show it in the mapping workflow; never discard it."),
    ("Late arrival", "Evaluate against the employee's assigned shift when Shift Management is enabled."),
], [2850, 6510])
doc.add_paragraph("The current Attendance Management UI should read processed biometric attendance while retaining existing leave and employee data. Multiple devices can contribute punches for the same employee and date.")

doc.add_heading("9. End-to-end test", level=1)
add_step("1. Connectivity", "Ping the machine IP. Verify Vertex ERP is listening on port 8082 and the firewall rule is enabled.")
add_step("2. Heartbeat", "Restart the machine. In Biometric Devices, verify Last Sync Time changes.")
add_step("3. Punch", "Use an enrolled fingerprint. Wait 5-30 seconds for the push request.")
add_step("4. Raw database check", "Confirm a new row exists in AttendanceLogs with the correct device, device user ID and punch time.")
add_step("5. Mapping check", "If EmployeeId is empty, create or correct EmployeeDeviceMapping.")
add_step("6. UI check", "Open Attendance Management, select the punch date, and run search. Verify employee, IN time, OUT time and status.")
add_step("7. Second punch", "Punch again later and confirm OUT time and working hours update to the latest punch.")

doc.add_heading("10. Troubleshooting", level=1)
add_table(["Symptom", "Most likely cause", "Corrective action"], [
    ("Machine IP becomes 0.0.0.0", "DHCP is ON but no DHCP server is available on a direct cable.", "Turn DHCP OFF and restore static 192.168.10.2/24."),
    ("Ping fails", "Wrong subnet, disconnected cable, disabled adapter or IP conflict.", "Check link lights; confirm laptop 192.168.10.1 and device 192.168.10.2."),
    ("Ping works, no Last Sync", "Wrong ADMS server/port, app not listening, firewall block or HTTPS mismatch.", "Use server 192.168.10.1:8082, HTTPS OFF; start ERP and allow TCP 8082."),
    ("HTTP 401 / unknown device", "Serial number is missing or does not match the registered device.", "Register the exact SN shown by the device request/device information."),
    ("Raw log exists, UI empty", "Device user is unmapped, wrong date filter or processing did not run.", "Map the user, process/backfill attendance, then search the punch date."),
    ("IN appears, OUT blank", "Only one valid punch exists for that date.", "Make a second punch; last punch becomes OUT."),
    ("Build error: file in use", "An existing Vertex ERP process is locking output files.", "Stop the running debug instance/process, rebuild, then start once."),
    ("Device sends to old server", "Cloud settings were not saved or the unit was not restarted.", "Save ADMS values and restart the machine."),
], [2300, 3020, 4040])

doc.add_heading("11. Production deployment checklist", level=1)
add_bullet("Host Vertex ERP as a continuously running Windows Service, IIS application, container or managed server process.")
add_bullet("Assign the server a fixed LAN IP or DNS name; avoid a changing laptop address.")
add_bullet("Restrict port 8082 to trusted device IPs/VPN. Use HTTPS only if the specific firmware configuration supports it end to end.")
add_bullet("Back up PostgreSQL and retain raw AttendanceLogs according to the attendance/audit policy.")
add_bullet("Use one unique record per device serial and maintain branch ownership for future multi-branch rollout.")
add_bullet("Monitor device heartbeat, last sync, rejected requests, unmapped users and database failures.")
add_bullet("Do not expose device communication keys, database passwords or public server addresses in screenshots or support tickets.")
add_bullet("Obtain employee notice/consent and follow applicable biometric and workplace privacy requirements.")

doc.add_heading("12. Acceptance criteria", level=1)
add_table(["Check", "Expected result"], [
    ("Device connectivity", "K40 Pro is reachable and Last Sync Time updates."),
    ("Raw capture", "Every new punch creates one durable AttendanceLogs record."),
    ("Employee mapping", "Device user resolves to the correct existing Employee."),
    ("Daily attendance", "First and last punch generate accurate IN/OUT times."),
    ("UI", "Attendance appears in the existing Attendance Management screen for the selected date."),
    ("Resilience", "Unmapped/duplicate/malformed data is logged without breaking unrelated ERP modules."),
    ("Scalability", "A second device can be registered without code or IP hardcoding."),
], [3000, 6360])

doc.add_heading("13. Support information to collect", level=1)
doc.add_paragraph("When escalating a problem, collect the following without sharing secrets:")
add_bullet("Device model, firmware version, push service version and serial number (share privately).")
add_bullet("Device Ethernet IP/subnet/gateway and ADMS server address/port/HTTPS setting.")
add_bullet("Laptop/server IP and whether port 8082 is listening.")
add_bullet("Exact punch time, Device User ID and employee name expected.")
add_bullet("Device Status/Last Sync Time and the related application log entries.")
add_bullet("Whether a raw AttendanceLogs row exists and whether EmployeeId is mapped.")

doc.add_heading("Quick reference", level=1)
add_table(["Item", "Current local value"], [
    ("Laptop/server Ethernet IP", "192.168.10.1"),
    ("K40 Pro Ethernet IP", "192.168.10.2"),
    ("Subnet mask", "255.255.255.0"),
    ("ADMS server", "192.168.10.1"),
    ("ADMS port", "8082"),
    ("HTTPS / Proxy", "OFF / OFF"),
    ("ADMS receive endpoint", "/iclock/cdata"),
    ("Device administration", "/BiometricDevices"),
    ("Attendance screen", "/Main/Attendence"),
], [3900, 5460])

# Keep tables readable and prevent rows from splitting where possible.
for table in doc.tables:
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit"); trPr.append(cant)

# Document properties
doc.core_properties.title = "ZKTeco K40 Pro Biometric Attendance Integration Guide"
doc.core_properties.subject = "Vertex ERP ADMS Push integration"
doc.core_properties.author = "Vertex ERP Implementation Team"
doc.core_properties.keywords = "ZKTeco, K40 Pro, ADMS, attendance, ASP.NET Core, PostgreSQL"

doc.save(OUT)
print(OUT)
